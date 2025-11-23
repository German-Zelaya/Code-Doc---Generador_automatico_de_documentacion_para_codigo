from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import markdown
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfgen import canvas

def create_docx(documented_code: str, filename: str, original_filename: str) -> BytesIO:
    """
    Genera un documento DOCX con el código documentado.
    """
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(10)
    
    # === PORTADA ===
    title = doc.add_heading('Documentación de Código', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(220, 38, 38)  # Rojo Doom
    
    # Información del archivo
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f'Archivo: {original_filename}')
    info_run.font.size = Pt(14)
    info_run.bold = True
    
    # Fecha de generación
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # === CÓDIGO DOCUMENTADO ===
    heading = doc.add_heading('Código Documentado', 1)
    heading_run = heading.runs[0]
    heading_run.font.color.rgb = RGBColor(234, 88, 12)  # Naranja
    
    # Agregar el código en un bloque
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(documented_code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)
    
    # Establecer fondo gris claro para el código
    shading_elm = code_para._element.get_or_add_pPr()
    
    doc.add_page_break()
    
    # === FOOTER ===
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        'Universidad Mayor Real y Pontificia de San Francisco Xavier de Chuquisaca\n'
        'Proyecto de Taller de Especialidad - SHC131\n'
        'Generador Automático de Documentación de Código'
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Guardar en BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def create_pdf_simple(documented_code: str, filename: str, original_filename: str) -> BytesIO:
    """
    Genera un PDF usando reportlab.
    """
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Estilos
        styles = getSampleStyleSheet()
        
        # Estilo de título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#dc2626',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Estilo de subtítulo
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor='#666666',
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        # Estilo de código
        code_style = ParagraphStyle(
            'Code',
            parent=styles['Code'],
            fontSize=8,
            leftIndent=20,
            fontName='Courier',
            spaceAfter=12
        )
        
        # Contenido del PDF
        story = []
        
        # Título
        story.append(Paragraph("Documentación de Código", title_style))
        story.append(Spacer(1, 12))
        
        # Información del archivo
        story.append(Paragraph(f"<b>Archivo:</b> {original_filename}", subtitle_style))
        story.append(Paragraph(f"<b>Generado:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", subtitle_style))
        story.append(Spacer(1, 30))
        
        # Sección de código
        story.append(Paragraph("<b>Código Documentado</b>", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        # Agregar código en bloques (para evitar problemas con código muy largo)
        code_lines = documented_code.split('\n')
        code_text = '<br/>'.join([line.replace('<', '&lt;').replace('>', '&gt;') for line in code_lines])
        
        # Usar Preformatted para el código
        story.append(Preformatted(documented_code, code_style))
        story.append(Spacer(1, 30))
        
        # Footer
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor='#999999',
            alignment=TA_CENTER
        )
        
        story.append(Spacer(1, 50))
        story.append(Paragraph(
            "Universidad Mayor Real y Pontificia de San Francisco Xavier de Chuquisaca<br/>"
            "Proyecto de Taller de Especialidad - SHC131<br/>"
            "Generador Automático de Documentación de Código",
            footer_style
        ))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        print(f"Error generando PDF: {e}")
        raise Exception(f"Error al generar PDF: {str(e)}")

def create_markdown_document(documented_code: str, filename: str, original_filename: str) -> str:
    """
    Genera un documento Markdown formateado.
    """
    markdown_content = f"""# Documentación de Código

**Archivo:** {original_filename}  
**Generado:** {datetime.now().strftime("%d/%m/%Y %H:%M")}

---

## Código Documentado

```python
{documented_code}
```

---

*Documentación generada automáticamente por Code Doc Generator*  
*Universidad Mayor Real y Pontificia de San Francisco Xavier de Chuquisaca*  
*Proyecto de Taller de Especialidad - SHC131*
"""
    
    return markdown_content