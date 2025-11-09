import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from ollama import chat
import os
from docx import Document
from fpdf import FPDF

# Importar markdown con fallback
try:
    import markdown as _markdown_lib
    markdown = _markdown_lib
except Exception:
    import re
    class _MarkdownFallback:
        @staticmethod
        def markdown(text):
            if text is None:
                return ''
            s = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            s = re.sub(r'(?m)^######\s*(.*)$', r'<h6>\1</h6>', s)
            s = re.sub(r'(?m)^#####\s*(.*)$', r'<h5>\1</h5>', s)
            s = re.sub(r'(?m)^####\s*(.*)$', r'<h4>\1</h4>', s)
            s = re.sub(r'(?m)^###\s*(.*)$', r'<h3>\1</h3>', s)
            s = re.sub(r'(?m)^##\s*(.*)$', r'<h2>\1</h2>', s)
            s = re.sub(r'(?m)^#\s*(.*)$', r'<h1>\1</h1>', s)
            s = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', s)
            s = re.sub(r'\*(.+?)\*', r'<em>\1</em>', s)
            s = re.sub(r'`(.+?)`', r'<code>\1</code>', s)
            s = s.replace('\n', '<br>\n')
            return s
    markdown = _MarkdownFallback()

class DocumentacionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🔥 DOOM Documentation Generator 🔥")
        self.root.configure(bg="#1c1c1c")
        
        # Variables de estado
        self.archivo_actual = None
        self.codigo = ""
        self.sugerencias = []
        self.sugerencias_aceptadas = {}
        self.indice_actual = 0
        
        self._configurar_ui()
        
    def _configurar_ui(self):
        # Estilos Doom
        self.estilo = {
            'bg': "#1c1c1c",
            'fg': "#ff3333",
            'btn_bg': "#555555",
            'btn_hover': "#ff3333",
            'font': ("Courier New", 12),
            'font_title': ("Courier New", 16, "bold")
        }
        
        # Título
        tk.Label(self.root, text="💀 DOOM Documentation Generator 💀",
                font=self.estilo['font_title'], fg=self.estilo['fg'],
                bg=self.estilo['bg']).pack(pady=10)
        
        # Frame superior
        top_frame = tk.Frame(self.root, bg=self.estilo['bg'])
        top_frame.pack(fill=tk.X, padx=10)
        
        # Botones principales
        self._crear_boton(top_frame, "Abrir Archivo", self.abrir_archivo).pack(side=tk.LEFT, padx=5)
        self._crear_boton(top_frame, "Nueva Sugerencia", self.regenerar_sugerencia).pack(side=tk.LEFT, padx=5)
        
        # Área de edición
        self.editor = scrolledtext.ScrolledText(
            self.root, width=80, height=20,
            bg="#111111", fg="#ff6666",
            font=self.estilo['font']
        )
        self.editor.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        # Frame inferior
        bottom_frame = tk.Frame(self.root, bg=self.estilo['bg'])
        bottom_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Botones de control
        self._crear_boton(bottom_frame, "← Anterior", self.anterior_sugerencia).pack(side=tk.LEFT, padx=5)
        self._crear_boton(bottom_frame, "Aceptar ✓", self.aceptar_sugerencia).pack(side=tk.LEFT, padx=5)
        self._crear_boton(bottom_frame, "Siguiente →", self.siguiente_sugerencia).pack(side=tk.LEFT, padx=5)

    def _crear_boton(self, parent, texto, comando):
        btn = tk.Button(parent, text=texto, command=comando,
                       font=self.estilo['font'],
                       bg=self.estilo['btn_bg'],
                       fg=self.estilo['fg'])
        btn.bind("<Enter>", lambda e: e.widget.config(bg=self.estilo['btn_hover']))
        btn.bind("<Leave>", lambda e: e.widget.config(bg=self.estilo['btn_bg']))
        return btn

    def abrir_archivo(self):
        """Abre un archivo de código y lo analiza."""
        ruta = filedialog.askopenfilename(
            filetypes=[("Archivos de código", "*.py;*.js;*.java;*.cpp;*.php")]
        )
        if ruta:
            self.archivo_actual = ruta
            with open(ruta, 'r', encoding='utf-8') as file:
                self.codigo = file.read()
            self.analizar_codigo()

    def analizar_codigo(self):
        """Analiza el código usando el modelo de IA y ast."""
        if not self.codigo:
            return
        
        self.editor.delete('1.0', tk.END)
        self.editor.insert(tk.END, "🔍 Analizando código...\n")
        
        try:
            import ast
            arbol = ast.parse(self.codigo)
            funciones = []
            
            for nodo in ast.walk(arbol):
                if isinstance(nodo, ast.FunctionDef):
                    inicio = nodo.lineno - 1
                    fin = nodo.end_lineno
                    lineas_funcion = self.codigo.split('\n')[inicio:fin]
                    codigo_funcion = '\n'.join(lineas_funcion)
                    
                    funciones.append({
                        'nombre': nodo.name,
                        'codigo': codigo_funcion,
                        'tiene_doc': ast.get_docstring(nodo) is not None
                    })
            
            if not funciones:
                self.editor.delete('1.0', tk.END)
                self.editor.insert(tk.END, "❌ No se encontraron funciones en el código.")
                return
            
            print(f"Funciones encontradas con ast: {len(funciones)}")
            self._analizar_funciones_con_ia(funciones)
            
        except Exception as e:
            print(f"Error en análisis ast: {str(e)}")
            messagebox.showerror("Error", f"Error al analizar el código: {str(e)}")

    def _analizar_funciones_con_ia(self, funciones):
        """Analiza las funciones encontradas usando el modelo de IA."""
        self.sugerencias = []
        
        for funcion in funciones:
            codigo = funcion['codigo']
            nombre = funcion['nombre']
            
            # Prompt mejorado para generar SIEMPRE en formato docstring
            prompt = f"""Eres un asistente experto en documentación de código. Tu tarea es generar ÚNICAMENTE un docstring en formato Python para la siguiente función.

CÓDIGO DE LA FUNCIÓN:
{codigo}

INSTRUCCIONES CRÍTICAS:
1. Genera SOLO el contenido del docstring (lo que va entre las triples comillas)
2. NO incluyas las triples comillas en tu respuesta
3. NO incluyas el nombre de la función ni "def"
4. NO incluyas comentarios adicionales
5. Usa el formato estándar de docstring Python:
   - Primera línea: Descripción breve de lo que hace la función
   - Args: (si hay parámetros) Lista de parámetros con descripción
   - Returns: (si retorna algo) Descripción del valor retornado

FORMATO ESPERADO:
Descripción breve y clara de la función.

Args:
    parametro1 (tipo): Descripción del parámetro
    parametro2 (tipo): Descripción del parámetro

Returns:
    tipo: Descripción del valor retornado

Responde SOLO con el contenido del docstring, sin formato extra."""

            try:
                print(f"\nAnalizando función: {nombre}")
                respuesta = chat(model="code-doc:latest", messages=[{"role": "user", "content": prompt}])
                contenido = respuesta["message"]["content"].strip()
                
                # Limpiar la respuesta de cualquier formato markdown o extra
                contenido = self._limpiar_docstring(contenido)
                
                self.sugerencias.append({
                    'funcion': nombre,
                    'docstring': contenido,
                    'codigo': codigo,
                    'estado': 'sin_documentar' if not funcion['tiene_doc'] else 'documentada'
                })
                print(f"Docstring generado para: {nombre}")
            
            except Exception as e:
                print(f"Error analizando función {nombre}: {str(e)}")
                continue
        
        if self.sugerencias:
            self.indice_actual = 0
            self.mostrar_sugerencia_actual()
        else:
            self.editor.delete('1.0', tk.END)
            self.editor.insert(tk.END, "❌ No se pudieron generar docstrings para las funciones.")

    def _limpiar_docstring(self, texto):
        """Limpia el texto del docstring removiendo formato markdown y extras."""
        # Remover bloques de código markdown
        texto = texto.replace('```python', '').replace('```', '')
        
        # Remover comillas triples si las incluyó por error
        texto = texto.replace('"""', '').replace("'''", '')
        
        # Remover líneas que contengan "def " o el nombre de función
        lineas = texto.split('\n')
        lineas_limpias = []
        for linea in lineas:
            if not linea.strip().startswith('def ') and 'FUNCIÓN:' not in linea.upper():
                lineas_limpias.append(linea)
        
        texto = '\n'.join(lineas_limpias).strip()
        
        return texto

    def mostrar_sugerencia_actual(self):
        """Muestra la sugerencia actual en formato docstring completo."""
        if not self.sugerencias:
            return
        
        sugerencia = self.sugerencias[self.indice_actual]
        self.editor.delete('1.0', tk.END)
        
        # Mostrar información de la función
        self.editor.insert(tk.END, f"🔧 Función: {sugerencia['funcion']}\n")
        self.editor.insert(tk.END, f"📊 Estado: {sugerencia['estado']}\n\n")
        
        # Mostrar el código original de la función
        self.editor.insert(tk.END, "📝 Código original:\n")
        self.editor.insert(tk.END, "=" * 60 + "\n")
        lineas_codigo = sugerencia['codigo'].split('\n')
        primera_linea = lineas_codigo[0] if lineas_codigo else ""
        self.editor.insert(tk.END, f"{primera_linea}\n")
        self.editor.insert(tk.END, "=" * 60 + "\n\n")
        
        # Mostrar el docstring sugerido en formato completo
        self.editor.insert(tk.END, "💡 Documentación sugerida (formato docstring):\n\n")
        self.editor.insert(tk.END, f'{primera_linea}\n')
        self.editor.insert(tk.END, '    """\n')
        
        # Indentar cada línea del docstring
        for linea in sugerencia['docstring'].split('\n'):
            if linea.strip():
                self.editor.insert(tk.END, f'    {linea}\n')
            else:
                self.editor.insert(tk.END, '\n')
        
        self.editor.insert(tk.END, '    """\n')
        
        # Mostrar contador
        total = len(self.sugerencias)
        self.editor.insert(tk.END, f"\n{'=' * 60}\n")
        self.editor.insert(tk.END, f"[📍 Función {self.indice_actual + 1} de {total}]")

    def regenerar_sugerencia(self):
        """Regenera una nueva sugerencia SIEMPRE en formato docstring."""
        if not self.sugerencias or self.indice_actual >= len(self.sugerencias):
            messagebox.showwarning("⚠️ Advertencia", "No hay función seleccionada para regenerar.")
            return
        
        sugerencia_actual = self.sugerencias[self.indice_actual]
        funcion_actual = sugerencia_actual["funcion"]
        codigo_funcion = sugerencia_actual["codigo"]
        
        # Prompt mejorado para regeneración - SIEMPRE formato docstring
        prompt = f"""Eres un asistente experto en documentación de código. Genera una NUEVA versión del docstring para esta función, con más detalles o desde otra perspectiva.

CÓDIGO DE LA FUNCIÓN:
{codigo_funcion}

DOCSTRING ANTERIOR:
{sugerencia_actual['docstring']}

INSTRUCCIONES CRÍTICAS:
1. Genera SOLO el contenido del docstring (sin las triples comillas)
2. Proporciona una versión más detallada o con enfoque diferente
3. Mantén el formato estándar de docstring Python
4. NO incluyas el nombre de la función ni "def"
5. NO incluyas comentarios adicionales fuera del docstring
6. Sé más específico en las descripciones de parámetros y retornos

FORMATO ESPERADO:
Descripción detallada de la función (puede ser más específica que antes).

Args:
    parametro (tipo): Descripción más detallada del parámetro

Returns:
    tipo: Descripción más detallada del retorno

Responde SOLO con el contenido del docstring mejorado."""

        try:
            self.editor.delete('1.0', tk.END)
            self.editor.insert(tk.END, f"🔄 Regenerando docstring para {funcion_actual}...\n")
            self.root.update()
            
            respuesta = chat(model="code-doc:latest", messages=[{"role": "user", "content": prompt}])
            nuevo_docstring = respuesta["message"]["content"].strip()
            
            # Limpiar el nuevo docstring
            nuevo_docstring = self._limpiar_docstring(nuevo_docstring)
            
            # Actualizar la sugerencia
            self.sugerencias[self.indice_actual]["docstring"] = nuevo_docstring
            
            # Mostrar el resultado
            self.mostrar_sugerencia_actual()
            
            print(f"Nueva versión de docstring generada para {funcion_actual}")
            messagebox.showinfo("✅ Éxito", f"Nueva documentación generada para {funcion_actual}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al regenerar docstring: {str(e)}")
            self.mostrar_sugerencia_actual()

    def anterior_sugerencia(self):
        """Muestra la sugerencia anterior."""
        if self.indice_actual > 0:
            self.indice_actual -= 1
            self.mostrar_sugerencia_actual()
        else:
            messagebox.showinfo("ℹ️ Info", "Ya estás en la primera función")

    def siguiente_sugerencia(self):
        """Muestra la siguiente sugerencia."""
        if self.indice_actual < len(self.sugerencias) - 1:
            self.indice_actual += 1
            self.mostrar_sugerencia_actual()
        else:
            messagebox.showinfo("ℹ️ Info", "Ya estás en la última función")

    def aceptar_sugerencia(self):
        """Acepta la sugerencia actual."""
        if not self.sugerencias or self.indice_actual >= len(self.sugerencias):
            return
        
        sugerencia_actual = self.sugerencias[self.indice_actual]
        funcion_actual = sugerencia_actual["funcion"]
        
        # Guardar el docstring aceptado
        self.sugerencias_aceptadas[funcion_actual] = {
            'docstring': sugerencia_actual['docstring'],
            'codigo': sugerencia_actual['codigo']
        }
        
        if self.indice_actual == len(self.sugerencias) - 1:
            # Si es la última sugerencia, mostrar vista previa
            messagebox.showinfo("✅ Completado", "¡Todas las funciones documentadas!\nGenerando vista previa...")
            self.mostrar_vista_previa()
        else:
            messagebox.showinfo("✅ Aceptado", f"Docstring aceptado para {funcion_actual}")
            self.siguiente_sugerencia()

    def mostrar_vista_previa(self):
        """Muestra una ventana de vista previa editable de la documentación final."""
        preview_window = tk.Toplevel(self.root)
        preview_window.title("🔥 Vista Previa - DOOM Documentation 🔥")
        preview_window.configure(bg="#1c1c1c")
        preview_window.geometry("900x700")
        
        # Título
        tk.Label(
            preview_window,
            text="💀 Documentación Final - Editable 💀",
            font=self.estilo['font_title'],
            fg=self.estilo['fg'],
            bg=self.estilo['bg']
        ).pack(pady=10)
        
        # Área de edición mejorada
        preview_text = scrolledtext.ScrolledText(
            preview_window,
            width=100,
            height=35,
            bg="#111111",
            fg="#ff6666",
            font=self.estilo['font'],
            insertbackground="#ff6666",
            insertwidth=2
        )
        preview_text.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        # Generar contenido en formato docstring
        contenido = f"# 📚 Documentación Generada - {os.path.basename(self.archivo_actual)}\n\n"
        contenido += "=" * 80 + "\n\n"
        
        for funcion, datos in self.sugerencias_aceptadas.items():
            contenido += f"## 🔧 Función: {funcion}\n\n"
            
            # Obtener la primera línea del código (def ...)
            primera_linea = datos['codigo'].split('\n')[0]
            
            contenido += "```python\n"
            contenido += f'{primera_linea}\n'
            contenido += '    """\n'
            
            # Agregar el docstring con indentación
            for linea in datos['docstring'].split('\n'):
                if linea.strip():
                    contenido += f'    {linea}\n'
                else:
                    contenido += '\n'
            
            contenido += '    """\n'
            contenido += "```\n\n"
            contenido += "-" * 80 + "\n\n"
        
        preview_text.insert('1.0', contenido)
        preview_text.focus_set()
        
        # Frame para botones
        btn_frame = tk.Frame(preview_window, bg="#1c1c1c")
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        def guardar_cambios():
            nuevo_contenido = preview_text.get('1.0', tk.END)
            messagebox.showinfo("✅ Éxito", "Cambios guardados. Ahora puedes exportar la documentación.")
            preview_window.destroy()
            self.mostrar_opciones_exportacion()
        
        # Botones con estilo DOOM
        self._crear_boton(btn_frame, "💾 Guardar y Exportar", guardar_cambios).pack(side=tk.LEFT, padx=5)
        self._crear_boton(btn_frame, "❌ Cancelar", preview_window.destroy).pack(side=tk.LEFT, padx=5)

    def mostrar_opciones_exportacion(self):
        """Muestra opciones para exportar la documentación."""
        export_window = tk.Toplevel(self.root)
        export_window.title("📤 Exportar Documentación")
        export_window.configure(bg="#1c1c1c")
        export_window.geometry("400x250")
        
        tk.Label(
            export_window,
            text="💀 Selecciona el formato de exportación 💀",
            font=self.estilo['font_title'],
            fg=self.estilo['fg'],
            bg=self.estilo['bg']
        ).pack(pady=20)
        
        btn_frame = tk.Frame(export_window, bg="#1c1c1c")
        btn_frame.pack(pady=20)
        
        self._crear_boton(btn_frame, "📄 Exportar como PDF", 
                         lambda: [self.exportar_doc('pdf'), export_window.destroy()]).pack(pady=10)
        self._crear_boton(btn_frame, "📝 Exportar como DOCX", 
                         lambda: [self.exportar_doc('docx'), export_window.destroy()]).pack(pady=10)
        self._crear_boton(btn_frame, "📋 Exportar como Markdown", 
                         lambda: [self.exportar_doc('md'), export_window.destroy()]).pack(pady=10)

    def exportar_doc(self, formato):
        """Exporta la documentación en el formato especificado."""
        if not self.sugerencias_aceptadas:
            messagebox.showwarning("⚠️ Advertencia", "No hay documentación para exportar.")
            return
        
        # Preparar contenido
        contenido = f"# Documentación de {os.path.basename(self.archivo_actual)}\n\n"
        
        for funcion, datos in self.sugerencias_aceptadas.items():
            primera_linea = datos['codigo'].split('\n')[0]
            contenido += f"## {funcion}\n\n"
            contenido += f"```python\n{primera_linea}\n"
            contenido += '    """\n'
            contenido += f"    {datos['docstring']}\n"
            contenido += '    """\n```\n\n'
        
        try:
            if formato == 'pdf':
                self._exportar_pdf(contenido)
            elif formato == 'docx':
                self._exportar_docx(contenido)
            elif formato == 'md':
                self._exportar_markdown(contenido)
            
            messagebox.showinfo("✅ Éxito", f"Documentación exportada en formato {formato.upper()}")
        except Exception as e:
            messagebox.showerror("Error", f"Error al exportar: {str(e)}")

    def _exportar_pdf(self, contenido):
        """Exporta el contenido a PDF."""
        ruta = filedialog.asksaveasfilename(defaultextension=".pdf")
        if ruta:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            texto = contenido.replace("#", "").replace("```python", "").replace("```", "")
            pdf.multi_cell(0, 10, txt=texto)
            pdf.output(ruta)

    def _exportar_docx(self, contenido):
        """Exporta el contenido a DOCX."""
        ruta = filedialog.asksaveasfilename(defaultextension=".docx")
        if ruta:
            doc = Document()
            doc.add_heading('Documentación Generada', 0)
            for linea in contenido.split('\n'):
                if linea.startswith('##'):
                    doc.add_heading(linea.replace('##', '').strip(), level=2)
                else:
                    doc.add_paragraph(linea)
            doc.save(ruta)

    def _exportar_markdown(self, contenido):
        """Exporta el contenido a Markdown."""
        ruta = filedialog.asksaveasfilename(defaultextension=".md")
        if ruta:
            with open(ruta, 'w', encoding='utf-8') as f:
                f.write(contenido)

if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentacionApp(root)
    root.mainloop()